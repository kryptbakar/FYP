"""Build the VYREX Software Requirements Specification (.docx).

Follows the provided FYP SRS template structure (IEEE-830 / ISO-IEC-IEEE-29148
lineage): cover, revision history, approval, TOC + list of figures/tables, then
Introduction, Overview, State of the Art, User/System Requirements, Functional
Requirements (with per-requirement traceability tables), Non-functional
Requirements, and the 4+1 architecture view model with UI design.

Every fact is grounded in the built system (this repo). Figures come from
figures/make_figures.py (run that first). Output: VYREX-SRS.docx next to this file.

Run:  python figures/make_figures.py && python make_srs.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).parent
FIG = HERE / "figures"
OUT = HERE / "VYREX-SRS.docx"

TEAL = RGBColor(0x1F, 0x7A, 0x6B)
INK = RGBColor(0x22, 0x26, 0x2B)
GREY = RGBColor(0x55, 0x5B, 0x62)

doc = Document()

# ---------------------------------------------------------------- base styles
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(11)
doc.styles["Title"].font.color.rgb = TEAL
for h, sz in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
    st = doc.styles[h]
    st.font.color.rgb = INK
    st.font.size = Pt(sz)


# ---------------------------------------------------------------- helpers
def para(text="", size=11, bold=False, italic=False, color=None, align=None,
         style=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def h1(t): return doc.add_heading(t, level=1)
def h2(t): return doc.add_heading(t, level=2)
def h3(t): return doc.add_heading(t, level=3)


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(" " + text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, bold=False, color=None, size=10, fill=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if fill:
        shade(cell, fill)


def table(rows, cols, header=None, widths=None, style="Light Grid Accent 1"):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = style
    except KeyError:
        t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header:
        for j, htext in enumerate(header):
            set_cell(t.rows[0].cells[j], htext, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                     fill="1F7A6B")
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    return t


def figure(fname, caption, width=6.3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG / fname), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption).italic = True
    cap.paragraph_format.space_after = Pt(10)


def field_toc():
    """Insert a Word TOC field (updates on open: right-click → Update Field, or F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and choose “Update Field” to build the table of contents."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, t, fldEnd):
        run._r.append(el)


def pagebreak():
    doc.add_page_break()


# ================================================================ COVER PAGE
for _ in range(2):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("VYREX"); r.bold = True; r.font.size = Pt(48); r.font.color.rgb = TEAL
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = s.add_run("An Air-Gapped Security Operations Center &\nVulnerability-Intelligence Platform")
rs.font.size = Pt(18); rs.font.color.rgb = INK
para("Software Requirements Specification (SRS)", size=15, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=2)
para("Version 1.0", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
for _ in range(2):
    doc.add_paragraph()
para("Bachelor of Science in Cyber Security — Final Year Design Project",
     size=12, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
para("Ghulam Ishaq Khan Institute of Engineering Sciences and Technology (GIKI)",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
para("Industry stakeholder: Punjab Information Technology Board (PITB)",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
for _ in range(1):
    doc.add_paragraph()

meta = table(4, 2, widths=[2.0, 4.0], style="Table Grid")
rows = [("Project Title", "VYREX — Air-Gapped SOC & Vulnerability-Intelligence Platform"),
        ("Date", "July 2026"),
        ("Supervisor", "____________________"),
        ("Group Members", "kryptbakar (Team Lead)  ·  ____________  ·  ____________")]
for i, (k, v) in enumerate(rows):
    set_cell(meta.rows[i].cells[0], k, bold=True, fill="EDF3F2")
    set_cell(meta.rows[i].cells[1], v)
pagebreak()

# ================================================================ REVISION / APPROVAL
para("Revision History", size=14, bold=True, color=INK)
rt = table(3, 3, header=["Version", "Date", "Comments"], widths=[1.2, 1.6, 3.6])
rev = [("1.00", "July 2026", "Initial SRS — baselined against the completed VYREX build (Phases 0–8)."),
       ("1.10", "July 2026", "Added evaluation, benchmarking, threat-model, and connector requirements.")]
for i, (v, d, c) in enumerate(rev):
    set_cell(rt.rows[i + 1].cells[0], v); set_cell(rt.rows[i + 1].cells[1], d)
    set_cell(rt.rows[i + 1].cells[2], c)
doc.add_paragraph()

para("Document Approval", size=14, bold=True, color=INK)
para("The following document has been accepted and approved by the following:", size=10.5, color=GREY)
at = table(4, 4, header=["Role", "Name", "Signature", "Date"], widths=[1.6, 2.0, 1.5, 1.2])
appr = [("Project Supervisor", "", "", ""), ("Co-Supervisor", "", "", ""),
        ("FYP Committee", "", "", "")]
for i, (role, n, sg, dt) in enumerate(appr):
    set_cell(at.rows[i + 1].cells[0], role, bold=True)
pagebreak()

# ================================================================ TOC + LISTS
para("Table of Contents", size=14, bold=True, color=INK)
field_toc()
pagebreak()
para("List of Figures", size=14, bold=True, color=INK)
figs = [
    "Figure 1 — SOC Analyst use cases", "Figure 2 — Administrator use cases",
    "Figure 3 — System / machine-actor use cases", "Figure 4 — Logical view: the four layers",
    "Figure 5 — Development view: repository modules", "Figure 6 — Process view: ingestion, scoring and response",
    "Figure 7 — Physical view: air-gapped deployment topology", "Figure 8 — Triage view (decision queue)",
    "Figure 9 — Finding detail (XAI / SHAP waterfall)", "Figure 10 — Cases view (incidents & audit)",
    "Figure 11 — Sensors & Fusion view (pipeline + tool grid)",
]
for f in figs:
    para(f, size=10.5, color=GREY, space_after=2)
doc.add_paragraph()
para("List of Tables", size=14, bold=True, color=INK)
tbls = ["Table 1 — Terms, acronyms and abbreviations", "Table 2 — Comparison with existing systems",
        "Table 3 — User classes and characteristics", "Table 4 — Functional requirements summary",
        "Table 5–24 — Functional requirement traceability tables", "Table 25 — Non-functional requirements",
        "Table 26 — Requirements-to-evaluation traceability"]
for t_ in tbls:
    para(t_, size=10.5, color=GREY, space_after=2)
pagebreak()

# ================================================================ 1. INTRODUCTION
h1("1  INTRODUCTION")
h2("1.1  Purpose")
para("This Software Requirements Specification (SRS) defines the functional and "
     "non-functional requirements of VYREX, an air-gapped Security Operations Center "
     "(SOC) and vulnerability-intelligence platform. It is intended for the project "
     "supervisors and FYP evaluation committee, the development team, and the industry "
     "stakeholder (PITB) as the agreed contract of what the system does. The document "
     "follows the IEEE-830 / ISO/IEC/IEEE-29148 recommended practice: it specifies "
     "*what* the system must do and the qualities it must exhibit, and traces each "
     "requirement to its rationale, source, and verification evidence, while leaving "
     "detailed *how* to the design and implementation.")
h2("1.2  Product Scope")
para("VYREX is a centralized SOC that runs fully on-premises and air-gapped: it ingests "
     "endpoint telemetry from a lightweight agent and normalized detections from ten "
     "integrated open-source security tools, enriches every vulnerability from a locally "
     "mirrored feed (CVSS + EPSS exploit probability + CISA KEV), and — its original "
     "contribution — fuses findings across tools and prioritizes them with an "
     "exploit-aware, explainable risk engine (a transparent composite score plus an "
     "XGBoost model with per-finding SHAP explanations). Analysts work a single "
     "risk-ranked queue and can execute signed, two-person-approved containment actions.")
para("In scope: telemetry collection; multi-tool ingestion and normalization; offline "
     "CVE/EPSS/KEV enrichment; CIS compliance evaluation with tamper-evident evidence; "
     "cross-tool fusion and risk scoring with explainability; an analyst console and "
     "Grafana dashboards; signed active response; and air-gapped deployment (Docker "
     "Compose and K3s). Out of scope: automated destructive response (containment only, "
     "always analyst-approved); a managed cloud/SaaS offering; and full Windows-agent "
     "feature parity in the MVP.")
para("Benefits: it lets budget-constrained government and defence-adjacent organizations "
     "— where foreign SaaS SIEMs are a non-starter because data cannot leave the premises "
     "— replace ten disjoint tool consoles with one prioritized, explained, and auditable "
     "workflow, with zero live internet dependency at runtime.")

h2("1.3  Definitions, Acronyms and Abbreviations")
terms = [
    ("SOC", "Security Operations Center — the team/platform that monitors and responds to threats."),
    ("Air-gapped", "A deployment with no runtime internet connectivity; feeds are mirrored offline."),
    ("CVE", "Common Vulnerabilities and Exposures — a public vulnerability identifier."),
    ("CVSS", "Common Vulnerability Scoring System — technical severity score (0–10)."),
    ("EPSS", "Exploit Prediction Scoring System (FIRST) — probability a CVE is exploited (0–1)."),
    ("KEV", "CISA Known Exploited Vulnerabilities catalog — CVEs exploited in the wild."),
    ("Finding", "A single security issue on an asset (a matched CVE, a detection, an IOC hit)."),
    ("Fusion", "Clustering findings that describe the same issue across independent tools."),
    ("Consensus", "A confidence weight rising with the number of independent tools that agree."),
    ("SHAP", "SHapley Additive exPlanations — per-feature attribution explaining a model score."),
    ("dedup_key", "A tool-independent key stamped on findings so fusion can cluster them."),
    ("mTLS", "Mutual TLS — both client and server authenticate with certificates."),
    ("MISP / OpenCTI / Sigma", "Threat-intel platforms and a detection-rule format, integrated as connectors."),
    ("ATT&CK", "MITRE ATT&CK — a knowledge base of adversary techniques (e.g. T1071 = C2)."),
    ("RBAC", "Role-Based Access Control — permissions by role (viewer / analyst / admin)."),
    ("SRS / FR / NFR", "This document / Functional Requirement / Non-functional Requirement."),
]
para("Table 1 — Terms, acronyms and abbreviations used in this document.", size=10, italic=True, color=GREY)
tt = table(len(terms) + 1, 2, header=["Term", "Description"], widths=[1.9, 4.4])
for i, (k, v) in enumerate(terms):
    set_cell(tt.rows[i + 1].cells[0], k, bold=True, size=9.5)
    set_cell(tt.rows[i + 1].cells[1], v, size=9.5)

h2("1.4  References")
for ref in [
    "IEEE Std 830-1998 / ISO/IEC/IEEE 29148:2018 — Recommended practice for Software Requirements Specifications.",
    "FIRST EPSS model documentation; CISA Known Exploited Vulnerabilities catalog; NIST NVD.",
    "MITRE ATT&CK framework; CIS Benchmarks (Debian 12).",
    "VYREX repository docs: ARCHITECTURE.md, DECISIONS.md (49 logged decisions), METHODOLOGY.md, "
    "THREAT-MODEL.md, BENCHMARKS.md, VALIDATION-ATTACK-SIM.md, PRODUCTION-DEPLOYMENT.md, CONNECTORS.md.",
    "Kruchten, P. (1995). Architectural Blueprints — The 4+1 View Model of Software Architecture.",
]:
    bullet(ref)
pagebreak()

# ================================================================ 2. OVERVIEW
h1("2  OVERVIEW")
h2("2.1  The Overall Description")
para("VYREX is organized in four layers: an endpoint agent, an ingestion & assessment "
     "pipeline, a data layer, and a presentation layer. Telemetry and tool detections "
     "flow up through a stateless, schema-validating edge into a durable message broker "
     "and asynchronous workers that enrich, fuse, and score findings; the results are "
     "presented to analysts as a ranked, explained queue and as dashboards. Exactly one "
     "component may reach the internet — a feed-sync job that mirrors vulnerability "
     "intelligence — and it is used offline everywhere else.")
h2("2.2  Product Perspective")
para("VYREX is a new, self-contained product that composes best-in-class open-source "
     "components (PostgreSQL, TimescaleDB, OpenSearch, NATS JetStream, Suricata, Zeek, "
     "Wazuh, Trivy, Nuclei, MISP, OpenCTI, Sigma, Falco) and contributes original value "
     "in the intelligence layer. It is not a plug-in to an existing SIEM; rather, tools "
     "attach to it as connectors (see §7). The design philosophy — integrate, don't "
     "reinvent; spend original effort on fusion, scoring and explainability — is recorded "
     "across 49 logged design decisions.")
figure("fig4_logical.png", "Figure 4 — Logical view: the four architectural layers of VYREX.")

h2("2.3  Product Functions")
for fn in [
    ("Collect", "endpoint telemetry (process, network, file-integrity, osquery host state) over mTLS."),
    ("Integrate", "ten OSS tools, normalizing every detection to one versioned telemetry envelope."),
    ("Enrich", "each CVE with CVSS + EPSS + KEV from a local mirror; evaluate CIS compliance with a hash-chained evidence log."),
    ("Fuse", "findings across tools by a tool-independent dedup key, deriving a consensus confidence weight."),
    ("Prioritize", "with a transparent 10-factor composite score and an XGBoost model, each explained per-finding by SHAP."),
    ("Respond", "via an Ed25519-signed command channel with two-person approval and a hash-chained audit trail (containment only)."),
    ("Present", "a real-time analyst console (triage, XAI detail, cases, compliance, fusion) plus Grafana dashboards."),
    ("Deploy", "fully air-gapped via Docker Compose or K3s, with a verified single-egress boundary and an offline installer bundle."),
]:
    bullet(fn[1], bold_lead=fn[0])

h2("2.4  User Characteristics")
para("Table 3 — User classes and characteristics.", size=10, italic=True, color=GREY)
users = [
    ("SOC Analyst", "Primary user", "Works the triage queue, inspects XAI detail, triages/labels findings, opens cases, requests containment. Security literate; not necessarily a developer."),
    ("Senior Analyst", "Approver", "All analyst abilities plus second-person approval of containment actions (admin role)."),
    ("SOC Administrator", "Operator", "Deploys and maintains the stack, manages users/RBAC, rolls out signed agents, retrains the model, sets defense policy, runs backups."),
    ("Staging-host Operator", "Feed custodian", "Runs the only internet-facing jobs (feed-sync / mirror-sync) on a connected DMZ host and carries updates across the gap."),
    ("Auditor / Assessor", "Read-only", "Verifies compliance posture and hash-chained audit integrity; viewer role."),
]
ut = table(len(users) + 1, 3, header=["User class", "Type", "Characteristics & responsibilities"], widths=[1.5, 1.2, 3.6])
for i, (a, b, c) in enumerate(users):
    set_cell(ut.rows[i + 1].cells[0], a, bold=True, size=9.5)
    set_cell(ut.rows[i + 1].cells[1], b, size=9.5)
    set_cell(ut.rows[i + 1].cells[2], c, size=9.5)

h2("2.5  Constraints")
for c in [
    ("Air-gap.", "No runtime internet access is permitted; only feed-sync (on a staging host) may egress, and the boundary is verifiable."),
    ("Linux-first server.", "The server stack runs on Linux (Ubuntu 22.04 / Rocky 9) via Docker/Kubernetes; agents are cross-platform."),
    ("On-premises only.", "All data and models remain on site; no third-party SaaS or telemetry egress."),
    ("Open-source licensing.", "GPL/AGPL components are used at arm's length (never vendored into the product); licenses are verified."),
    ("Explainability.", "Every risk score must be defensible: composite contributions and SHAP attributions are mandatory."),
    ("Two-person control.", "No destructive action; containment requires two distinct approvals and is cryptographically signed and audited."),
]:
    bullet(c[1], bold_lead=c[0])

h2("2.6  Assumptions and Dependencies")
for a in [
    "A connected staging host of the same CPU architecture is available periodically to refresh the offline feed/tool mirror.",
    "Network sensors (Suricata/Zeek) are fed by a SPAN/mirror port or TAP — a physical prerequisite that cannot be met in software.",
    "Endpoints can run the signed Go agent and are issued mTLS client certificates from the site PKI (Vault in K3s).",
    "Docker Engine + Compose (pilot) or a K3s cluster (production) with sufficient CPU/RAM/SSD per the sizing tables is provided.",
    "The synthetic training prior is replaced by real analyst-feedback labels as a deployment accrues them (retraining loop).",
]:
    bullet(a)
pagebreak()

# ================================================================ 3. STATE OF THE ART
h1("3  STATE OF THE ART")
h2("3.1  Literature Review")
para("Vulnerability management has shifted from severity-only (CVSS) triage toward "
     "exploit-aware prioritization. Research and practice (FIRST's EPSS, CISA's KEV "
     "catalog) show that CVSS alone over-prioritizes: most high-CVSS CVEs are never "
     "exploited, while a small, EPSS/KEV-identifiable subset drives real risk. Modern "
     "detection stacks are also multi-tool, producing overlapping, un-correlated alerts; "
     "the analyst burden is correlation and prioritization, not detection. Finally, "
     "explainable AI (SHAP / Shapley additive attributions) has become the accepted way "
     "to make ML-based scoring auditable — essential in a security context where an "
     "analyst must justify why a finding was escalated. VYREX synthesizes these threads: "
     "exploit-aware scoring, cross-tool fusion, and per-finding explanation, under a hard "
     "air-gap constraint that rules out mainstream cloud SIEMs.")
h2("3.2  Existing Systems")
para("Table 2 — Comparison with existing systems.", size=10, italic=True, color=GREY)
comp = [
    ("Capability", "Wazuh (alone)", "Elastic Security", "Splunk ES", "VYREX"),
    ("Fully air-gapped, no telemetry egress", "Yes", "Self-managed only", "Self-managed only", "Yes — verified"),
    ("Host+network+vuln+intel in one pipeline", "Host-centric", "Broad (paid)", "Broad (paid)", "Yes — 10 OSS tools"),
    ("Cross-tool consensus as a ranking signal", "No", "Manual rules", "Manual searches", "Yes — automatic"),
    ("Exploit-aware scoring (EPSS+KEV) built in", "Partial (CVSS)", "Via integrations", "Via apps", "Yes — native"),
    ("Per-finding ML explanation (SHAP)", "No", "Limited", "Limited", "Yes — TreeSHAP"),
    ("Signed, two-person active response", "Limited", "Via connectors", "Via SOAR (paid)", "Yes — Ed25519"),
    ("Cost model", "OSS", "Paid tiers", "Licensed", "OSS + on-prem"),
]
ct = table(len(comp), 5, widths=[2.3, 1.15, 1.15, 1.05, 1.2])
for j, htext in enumerate(comp[0]):
    set_cell(ct.rows[0].cells[j], htext, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), fill="1F7A6B", size=9)
for i, row in enumerate(comp[1:], start=1):
    for j, v in enumerate(row):
        set_cell(ct.rows[i].cells[j], v, size=9, bold=(j == 4))
para("VYREX does not aim to out-feature Splunk on connector breadth or scale-out "
     "maturity. Its thesis is that for the air-gapped, budget-constrained, "
     "prioritization-overwhelmed SOC, a fusion-and-explanation layer over best-in-class "
     "OSS beats both a single heavyweight vendor and ten disjoint consoles.", size=10.5, color=GREY, italic=True)
pagebreak()

# ================================================================ 4. USER/SYSTEM REQ
h1("4  USER / SYSTEM REQUIREMENTS")
para("This section captures the requirements from the users' and the system's viewpoint, "
     "then the external interfaces. Actors and their goals are shown in the use-case "
     "diagrams (Figures 1–3); the functional requirements of §5 realize these use cases.")
figure("fig1_usecases_analyst.png", "Figure 1 — SOC Analyst and Senior Analyst use cases.")
figure("fig2_usecases_admin.png", "Figure 2 — Administrator and Staging-host Operator use cases.")
figure("fig3_usecases_system.png", "Figure 3 — System / machine-actor use cases (agent, pipeline, tools).")

h2("4.1  External Interface Requirements")
h3("4.1.1  User Interfaces")
for u in [
    "A dependency-free single-page analyst console served over HTTPS with five views: Triage, Finding Detail (XAI), Compliance, Cases, and Sensors & Fusion (see Figures 8–11).",
    "A conclusion-first, instrument-grade design language (charcoal canvas, single teal accent on interactive chrome only), WCAG-AA severity encoding by shape and label (not color alone), and full keyboard navigation.",
    "Grafana dashboards for metrics, trends and heatmaps; Swagger UI for the REST API.",
]:
    bullet(u)
h3("4.1.2  Hardware Interfaces")
for u in [
    "Endpoint hosts: the Go agent reads process/network activity, file-integrity events and osquery host state; resource-capped (configurable CPU/memory).",
    "Network sensor host: a NIC in promiscuous mode fed by a SPAN/mirror port or TAP for Suricata/Zeek.",
    "Server: x86-64 Linux host(s) with SSD/NVMe storage per the sizing tables (pilot 8c/32GB/1TB; production 3+ K3s nodes).",
]:
    bullet(u)
h3("4.1.3  Software Interfaces")
for u in [
    "Data stores: PostgreSQL (transactional state), TimescaleDB (telemetry hypertable), OpenSearch (search).",
    "Broker: NATS JetStream (durable subjects telemetry.v1.<kind>).",
    "Integrated tools as connectors: Suricata, Zeek, Wazuh, Falco, Trivy, Nuclei, MISP, OpenCTI, Sigma.",
    "Local feed mirror: NVD, EPSS, CISA KEV; tool feeds (Nuclei templates, Trivy DB, Sigma rules, ET Open).",
    "Identity: Keycloak (OIDC/RBAC) and HashiCorp Vault (secrets, PKI) in the K3s deployment.",
]:
    bullet(u)
h3("4.1.4  Communication Interfaces")
for u in [
    "Agent → ingest-edge: mutual TLS (TLS 1.2+) on port 8443 with a bearer token; the agent_id is cross-checked against the client-cert CN.",
    "Server → agent: an Ed25519-signed command channel (nonce + expiry) for containment; the agent verifies fail-closed before executing.",
    "Analyst → console/API: HTTPS with OIDC (Keycloak) or local session tokens; RBAC enforced on every non-public route.",
    "feed-sync → internet: the single permitted egress, NetworkPolicy-enforced and provable via `make airgap-verify`.",
]:
    bullet(u)
pagebreak()

# ================================================================ 5. FUNCTIONAL REQ
h1("5  FUNCTIONAL REQUIREMENTS")
para("The functional requirements are listed below and then specified individually with "
     "traceability information (each in its own table, per the template). Requirement IDs "
     "are of the form FR-n; use cases are UC-n as shown in Figures 1–3.")

FRS = [
    ("FR-01", "Telemetry collection", "High", "UC-Collect",
     "The endpoint agent shall collect process events, network flows, file-integrity events and osquery host-state and ship them as schema-valid telemetry envelopes.",
     "Endpoints must be observable to detect and enrich threats.", "PITB requirement; ARCHITECTURE §Agent"),
    ("FR-02", "Secure ingestion", "High", "UC-Ship / UC-Ingest",
     "The system shall accept telemetry only over mutual TLS with a valid bearer token, validate each envelope against schema v1, and reject non-conforming or mismatched (agent_id≠cert CN) events.",
     "Prevents spoofed or malformed telemetry from poisoning the pipeline.", "THREAT-MODEL TB1"),
    ("FR-03", "Durable queueing & storage", "High", "UC-Ingest",
     "The system shall enqueue validated telemetry on a durable broker (JetStream) and persist it to TimescaleDB and OpenSearch idempotently (event_id as key).",
     "Decouples ingest from processing and protects stores via back-pressure.", "ARCHITECTURE §Ingestion"),
    ("FR-04", "Multi-tool integration", "High", "UC-Sensors",
     "The system shall integrate ten OSS tools (Suricata, Zeek, Wazuh, Falco, Trivy, Nuclei, MISP, OpenCTI, Sigma, and the agent), normalizing every detection into findings or envelopes.",
     "One normalized pipeline replaces ten disjoint consoles.", "Product scope; CONNECTORS.md"),
    ("FR-05", "Offline CVE enrichment", "High", "UC-Enrich",
     "The system shall match discovered packages/CVEs and attach CVSS, EPSS and KEV from the local mirror, with no live internet call at runtime.",
     "Exploit-aware context is the basis of prioritization; air-gap forbids live calls.", "Literature (EPSS/KEV); AIRGAP.md"),
    ("FR-06", "CIS compliance evaluation", "Medium", "UC-Comply",
     "The system shall evaluate assets against CIS-benchmark-style rules producing pass/fail/partial/not-applicable, recording each result with a hash-chained evidence entry.",
     "Regulated buyers require auditable hardening posture.", "PITB compliance need"),
    ("FR-07", "Cross-tool fusion", "High", "UC-Fuse",
     "The system shall cluster findings sharing a tool-independent dedup_key into one issue and compute a consensus weight rising with the number of distinct corroborating tools.",
     "Independent corroboration is a strong, original prioritization signal.", "FUSION.md; original contribution"),
    ("FR-08", "Composite risk scoring", "High", "UC-Score",
     "The system shall compute a transparent composite risk score (0–100) as a weighted sum of ten normalized factors, recording each factor's point contribution.",
     "A defensible, fully explainable baseline score.", "scoring.py; METHODOLOGY §FR4"),
    ("FR-09", "ML risk scoring", "High", "UC-Score",
     "The system shall additionally score each finding with a trained XGBoost model capturing non-linear factor interactions.",
     "Learned interactions (KEV×EPSS, exposure×CVSS) improve ranking over the linear score.", "train.py; evaluate.py"),
    ("FR-10", "Per-finding explainability", "High", "UC-Detail",
     "The system shall produce, for every ML score, an exact TreeSHAP attribution reconciling base + Σ contributions ≈ final, surfaced as a waterfall in the console.",
     "Every escalation must be auditable and defensible.", "explain.py; CONSOLE §SHAP"),
    ("FR-11", "Threat-intel & ATT&CK enrichment", "Medium", "UC-Intel",
     "The system shall enrich findings with MISP IOC matches, MITRE ATT&CK technique mapping, and Sigma detections.",
     "Real-world activity and adversary context sharpen prioritization.", "intel-enricher; features.py"),
    ("FR-12", "Risk-ranked triage queue", "High", "UC-Triage",
     "The system shall present findings as a decision queue ranked by composite risk, filterable by domain, severity, source tool and KEV.",
     "Analysts must work the most important findings first.", "CONSOLE §Triage"),
    ("FR-13", "Analyst feedback loop", "Medium", "UC-Feedback",
     "The system shall let an analyst assign a priority label to a finding and shall fold sanitized feedback into model retraining at bounded influence.",
     "Turns triage into training signal while resisting label poisoning.", "feedback.py; THREAT-MODEL ML"),
    ("FR-14", "Incident case management", "Medium", "UC-Case",
     "The system shall manage incident cases with status and SLA and a hash-chained audit timeline of all actions.",
     "Investigations need tracking and a tamper-evident record.", "CONSOLE §Cases"),
    ("FR-15", "Signed active response", "High", "UC-Request / UC-Exec",
     "The system shall issue containment commands only as Ed25519-signed messages (nonce+expiry) that the agent verifies fail-closed before executing; commands are a fixed containment allow-list, never arbitrary shell.",
     "A hijacked response channel is RCE-equivalent; it must be authenticated and bounded.", "D-028; THREAT-MODEL TB5"),
    ("FR-16", "Two-person approval", "High", "UC-Approve",
     "The system shall require two distinct approver identities before any containment action is authorized, enforced server-side.",
     "Prevents unilateral or mistaken destructive action.", "D-027/D-028"),
    ("FR-17", "Authentication & RBAC", "High", "UC-Users",
     "The system shall require an authenticated principal (OIDC header or session token) on every non-public route and enforce role permissions (viewer read-only, analyst read+write, admin for response/defense); authentication is mandatory in production.",
     "Protects the crown-jewel telemetry and the response channel.", "auth_guard.py; THREAT-MODEL TB2"),
    ("FR-18", "Air-gap enforcement & verification", "High", "UC-Airgap",
     "The system shall confine internet egress to feed-sync alone and provide a harness that proves the runtime boundary is sealed.",
     "The core promise to an air-gapped buyer must be provable, not asserted.", "airgap-verify; NFR air-gap"),
    ("FR-19", "Offline installation & updates", "Medium", "UC-Bundle",
     "The system shall build a checksummed offline bundle (images + feed mirror + config) on a connected host and install it inside the air gap with fail-closed integrity verification.",
     "Software and feed updates must cross the gap safely and repeatably.", "bundle.sh/install.sh; PRODUCTION-DEPLOYMENT §3"),
    ("FR-20", "Dashboards & search", "Medium", "UC-Dash / UC-Hunt",
     "The system shall provide Grafana dashboards (metrics/trends/heatmaps) and global search over assets, CVEs and IOCs across the stores.",
     "Situational awareness and hunting complement the triage queue.", "grafana/; search router"),
]

para("Table 4 — Functional requirements summary.", size=10, italic=True, color=GREY)
st = table(len(FRS) + 1, 4, header=["ID", "Requirement", "Priority", "Use case"], widths=[0.8, 3.6, 0.9, 1.4])
for i, fr in enumerate(FRS):
    set_cell(st.rows[i + 1].cells[0], fr[0], bold=True, size=9)
    set_cell(st.rows[i + 1].cells[1], fr[1], size=9)
    set_cell(st.rows[i + 1].cells[2], fr[2], size=9)
    set_cell(st.rows[i + 1].cells[3], fr[3], size=9)
pagebreak()

h2("5.1  Functional Requirements with Traceability Information")
para("Each requirement is specified below in its own traceability table.", size=10.5, color=GREY)
for idx, (rid, name, prio, uc, desc, rationale, source) in enumerate(FRS):
    trace = table(6, 2, widths=[1.7, 4.6], style="Light List Accent 1")
    fields = [("Requirement ID", f"{rid}  —  {name}"),
              ("Status", "Baselined  ·  Priority: " + prio),
              ("Description", desc),
              ("Rationale", rationale),
              ("Source", source),
              ("Use case / Verification", f"{uc}  ·  Verified by: unit tests / evaluation harness / e2e smoke (see §6.2, Table 26)")]
    for r_, (k, v) in enumerate(fields):
        set_cell(trace.rows[r_].cells[0], k, bold=True, size=9.5, fill="EDF3F2")
        set_cell(trace.rows[r_].cells[1], v, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
pagebreak()

# ================================================================ 6. NON-FUNCTIONAL
h1("6  NON-FUNCTIONAL REQUIREMENTS & SOFTWARE SYSTEM ATTRIBUTES")
h2("6.1  Performance Requirements")
para("Performance targets are measured with the in-repo benchmark harness "
     "(BENCHMARKS.md: ingest throughput, end-to-end latency, API load, footprint). "
     "The values below are the target envelope for a single-node pilot; production K3s "
     "scales stateless workers horizontally.")
NFRS = [
    ("NFR-01", "Performance", "Ingest throughput", "Sustain the pilot's measured events/sec through edge→queue→workers→stores without unbounded lag (ramp 5k→200k envelopes)."),
    ("NFR-02", "Performance", "End-to-end latency", "Median (p50) telemetry-to-queryable latency within the documented target; p95/p99 reported."),
    ("NFR-03", "Performance", "API responsiveness", "Console hot-path endpoints meet p95 < 400 ms at 20 virtual users, error rate < 1% (k6 gate)."),
    ("NFR-04", "Security", "Confidentiality & authN", "mTLS ingestion; mandatory OIDC/session auth + RBAC on every non-public route in production."),
    ("NFR-05", "Security", "Integrity / tamper-evidence", "Compliance evidence and response audit are hash-chained and verifiable; agent supply chain is cosign-signed."),
    ("NFR-06", "Security", "Air-gap", "Zero runtime egress except feed-sync; boundary provable via airgap-verify."),
    ("NFR-07", "Reliability", "Durability & back-pressure", "No telemetry loss under sustained load; broker back-pressure protects the stores; idempotent writes."),
    ("NFR-08", "Availability", "Recoverability", "Scheduled backup/restore of all stores (Velero/backup); tested restore procedure."),
    ("NFR-09", "Explainability", "Auditable scoring", "Every score exposes composite factor contributions and exact SHAP attributions; fixed seeds make ML reproducible."),
    ("NFR-10", "Usability", "Accessibility", "WCAG-AA severity encoding by shape+label (not color alone); full keyboard navigation."),
    ("NFR-11", "Maintainability", "Tested & CI-gated", "Automated tests (91) plus Go vet, image scan, and k3d + Compose e2e smokes run in CI on every change."),
    ("NFR-12", "Portability", "Deployment footprints", "Runs as single-node Docker Compose (pilot) and multi-node K3s via a Helm chart (production)."),
    ("NFR-13", "Scalability", "Horizontal workers", "Stateless ingest and workers scale out in K3s to grow ingest capacity for larger estates."),
    ("NFR-14", "Extensibility", "Pluggable connectors", "New tools attach via the documented connector contract (envelope or findings + dedup_key) with no core change."),
]
para("Table 25 — Non-functional requirements and software system attributes.", size=10, italic=True, color=GREY)
nt = table(len(NFRS) + 1, 4, header=["ID", "Attribute", "Name", "Requirement"], widths=[0.8, 1.3, 1.4, 3.2])
for i, (a, b, c, d) in enumerate(NFRS):
    set_cell(nt.rows[i + 1].cells[0], a, bold=True, size=9)
    set_cell(nt.rows[i + 1].cells[1], b, size=9)
    set_cell(nt.rows[i + 1].cells[2], c, size=9)
    set_cell(nt.rows[i + 1].cells[3], d, size=9)

h2("6.2  Verification & Requirements Traceability")
para("Table 26 — Requirements-to-evaluation traceability (how each requirement is verified).", size=10, italic=True, color=GREY)
verif = [
    ("FR-07/08/09/10 (fusion, scoring, ML, XAI)", "ml/evaluate.py + eval_fusion.py + attack_scenario.py; unit tests (fusion, scoring, features, dataset, evaluate, feedback, attack-scenario)."),
    ("FR-05/06 (enrichment, compliance)", "services/enrichment/tests (version→CVE matching; compliance pass/fail/NA)."),
    ("FR-17 (authN + RBAC)", "services/api/tests/test_auth_guard.py (401 unauth, RBAC 403s, production-forces-auth)."),
    ("FR-01..05,12,20 (end-to-end pipeline)", "deploy/smoke/compose-smoke.sh: ingest→assess→score→assert findings carry a risk_score."),
    ("NFR-01/02/03 (performance)", "tools/fake-producer/e2e_latency.py, make bench-ingest, tools/load/k6-api.js."),
    ("NFR-06 (air-gap)", "tools/airgap/verify-egress.sh (make airgap-verify)."),
    ("NFR-11 (CI gates)", ".github/workflows/ci.yml: pytest (91), go vet+build, Trivy scan, k3d smoke, compose e2e smoke."),
]
vt = table(len(verif) + 1, 2, header=["Requirement(s)", "Verification evidence"], widths=[2.6, 3.7])
for i, (a, b) in enumerate(verif):
    set_cell(vt.rows[i + 1].cells[0], a, bold=True, size=9)
    set_cell(vt.rows[i + 1].cells[1], b, size=9)
pagebreak()

# ================================================================ 7. DESIGN / 4+1
h1("7  PROJECT DESIGN / ARCHITECTURE")
para("The architecture is documented with Kruchten's 4+1 view model. The Use-Case view "
     "(Figures 1–3, §4) is the “+1” that ties the others together; the Logical, "
     "Development, Process and Physical views follow.")
h2("7.1  Use-Case View")
para("Actors and goals are shown in Figures 1–3: SOC Analyst and Senior Analyst (triage, "
     "XAI inspection, feedback, cases, containment request and approval); Administrator "
     "and Staging-host Operator (deployment, users/RBAC, agent rollout, feed mirroring, "
     "offline bundle, retraining, air-gap verification); and machine actors — the "
     "Endpoint Agent, the Pipeline Services and the Integrated Tools.")
h2("7.2  Logical View")
para("Figure 4 shows the four cooperating layers and the responsibilities and key "
     "collaborations of each. The intelligence layer (enrichment + fusion + scoring + "
     "explanation) is where the original contribution lives; everything else composes "
     "proven OSS.")
h2("7.3  Development View")
figure("fig5_development.png", "Figure 5 — Development view: repository modules and their roles.")
para("VYREX is a single repository whose modules map cleanly onto the layers: agent/, "
     "services/ (api, ingest-edge, workers, enrichment, feed-sync, bridges), ml/ "
     "(fusion, scoring, training, evaluation), web/console/, schema/, deploy/ (Helm, "
     "Vault, Keycloak, smokes), tools/ (producer, load, airgap), and docs/. Docker "
     "Compose files and a Makefile drive every module; CI gates the whole.")
h2("7.4  Process View")
figure("fig6_process.png", "Figure 6 — Process view: the ingestion, scoring and signed-response flows.")
para("Three runtime flows cooperate: (1) stateless telemetry ingestion into a durable "
     "queue; (2) the assessment & scoring loop (enrich → fuse → composite + ML + SHAP → "
     "rank); and (3) analyst-controlled, two-person, Ed25519-signed containment with a "
     "hash-chained audit.")
h2("7.5  Physical View")
figure("fig7_physical.png", "Figure 7 — Physical view: air-gapped deployment topology.")
para("Endpoints, sensor host, analyst workstations, the server cluster and data stores "
     "all sit inside the sealed boundary. Only a DMZ staging host reaches the internet to "
     "mirror feeds; updates cross the gap on verified removable media via the offline "
     "installer bundle. The single-egress invariant is enforced (NetworkPolicy) and "
     "proven (airgap-verify).")

h2("7.6  User Interface Design")
para("The console realizes the presentation requirements in a conclusion-first, "
     "instrument-grade design language (charcoal canvas, one teal accent used only on "
     "interactive chrome, severity by shape and label for WCAG-AA). The four wireframes "
     "below show the primary screens.")
figure("fig8_ui_triage.png", "Figure 8 — Triage view: the risk-ranked decision queue (each row = a conclusion + one action).")
figure("fig9_ui_detail.png", "Figure 9 — Finding detail: the XAI drawer with composite + ML scores, the SHAP waterfall, the containment gate and evidence provenance.")
figure("fig10_ui_cases.png", "Figure 10 — Cases view: incidents with SLA and a hash-chained audit timeline.")
figure("fig11_ui_fusion.png", "Figure 11 — Sensors & Fusion view: the live pipeline and the integrated-tool grid.")

pagebreak()
h1("8  CONCLUSION")
para("This SRS specifies VYREX as an air-gapped SOC whose distinguishing value is the "
     "intelligence layer: exploit-aware, cross-tool, explainable prioritization over "
     "best-in-class open-source detection. The twenty functional and fourteen "
     "non-functional requirements are each traceable to a rationale, a source, and — "
     "importantly — concrete verification evidence already present in the build (unit and "
     "component tests, evaluation and benchmark harnesses, an end-to-end smoke, and "
     "air-gap verification). The system is implemented and its claims are measured rather "
     "than asserted, satisfying both the academic bar for a final-year design project and "
     "the due-diligence bar for a deployable product.")

# ---------------------------------------------------------------- page numbers
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("VYREX — Software Requirements Specification v1.0        Page ")
run.font.size = Pt(8); run.font.color.rgb = GREY
fld1 = OxmlElement("w:fldSimple"); fld1.set(qn("w:instr"), "PAGE")
fp._p.append(fld1)

doc.save(str(OUT))
print("SRS written to", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
